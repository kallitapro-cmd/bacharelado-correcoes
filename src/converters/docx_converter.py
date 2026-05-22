"""Conversor DOCX -> texto. Story 1.7 — python-docx (extracao nativa)."""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject

MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


def convert_docx(file_path: str | Path) -> dict[str, object]:
    """Converte um arquivo DOCX em texto.

    Extrai parágrafos via doc.paragraphs e células de tabela via doc.tables,
    concatenando parágrafos primeiro e depois tabelas, separados por '\\n'.

    Args:
        file_path: Caminho para o arquivo DOCX.

    Returns:
        dict com campos:
        - texto (str): texto extraído (parágrafos + tabelas)
        - paragrafos (int): número de parágrafos extraídos
        - metodo_extracao (str): "nativo" em sucesso; "nenhum" em falha
        - flags (list[str]): vazio = sucesso; possíveis valores:
            "formato_nao_suportado" (arquivo .doc legado),
            "arquivo_ilegivel" (corrompido ou inacessível)

    Raises:
        ValueError: se o arquivo exceder MAX_FILE_BYTES (10 MB).
    """
    file_path = Path(file_path)

    # R1 — rejeitar .doc legado antes de tentar abrir (python-docx não suporta)
    nome_lower = file_path.name.lower()
    if nome_lower.endswith(".doc") and not nome_lower.endswith(".docx"):
        return _falha("formato_nao_suportado")

    # Checar tamanho antes de abrir (AC3)
    try:
        tamanho = file_path.stat().st_size
    except OSError:
        return _falha("arquivo_ilegivel")

    if tamanho > MAX_FILE_BYTES:
        mb = tamanho // (1024 * 1024)
        raise ValueError(
            f"Arquivo muito grande: {mb}MB. Limite: {MAX_FILE_BYTES // (1024 * 1024)}MB."
        )

    try:
        doc = Document(str(file_path))
    except PackageNotFoundError:
        return _falha("arquivo_ilegivel")
    except Exception:
        # M1 — captura ampla para qualquer falha de leitura/parsing do XML
        return _falha("arquivo_ilegivel")

    try:
        paragrafos_texto = _extrair_paragrafos(doc)
        tabelas_texto = _extrair_tabelas(doc)

        partes: list[str] = []
        if paragrafos_texto:
            partes.append(paragrafos_texto)
        if tabelas_texto:
            partes.append(tabelas_texto)
        texto_final = "\n".join(partes)

        return {
            "texto": _sanitizar(texto_final),
            "paragrafos": len(doc.paragraphs),
            "metodo_extracao": "nativo",
            "flags": [],
        }
    except Exception:
        # M2 — qualquer falha de extração tardia também retorna flag (sem propagar)
        return _falha("arquivo_ilegivel")


def _extrair_paragrafos(doc: DocumentObject) -> str:
    """Extrai texto de todos os parágrafos do documento."""
    linhas: list[str] = []
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        if texto:
            linhas.append(texto)
    return "\n".join(linhas)


def _extrair_tabelas(doc: DocumentObject) -> str:
    """Extrai texto das células de todas as tabelas (ordem: tabela > linha > célula)."""
    partes: list[str] = []
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas: list[str] = []
            for celula in linha.cells:
                texto_celula = celula.text.strip()
                if texto_celula:
                    celulas.append(texto_celula)
            if celulas:
                partes.append(" ".join(celulas))
    return "\n".join(partes)


def _sanitizar(texto: str) -> str:
    """Remove caracteres de controle e normaliza encoding UTF-8."""
    texto = unicodedata.normalize("NFC", texto)
    texto = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", texto)
    return texto.strip()


def _falha(flag: str) -> dict[str, object]:
    """Retorna estrutura padrão de falha com flag indicando o motivo."""
    return {
        "texto": "",
        "paragrafos": 0,
        "metodo_extracao": "nenhum",
        "flags": [flag],
    }
